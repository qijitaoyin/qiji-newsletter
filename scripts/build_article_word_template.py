from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "導引電子報文章撰寫範本_精簡版.docx"
HERO_IMAGE = ROOT / "public" / "assets" / "articles" / "202403" / "202403-img-003.jpg"
BODY_IMAGE = ROOT / "public" / "assets" / "articles" / "202403" / "202403-img-005.jpg"

INK = RGBColor(35, 38, 34)
CONTROL = RGBColor(180, 45, 45)
MUTED = RGBColor(105, 108, 101)


def set_font(run, size=11, bold=None, color=INK, name="Microsoft JhengHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_spacing(paragraph, before=0, after=6, line=1.4):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(doc, text, size=11, bold=None, before=0, after=6, line=1.4):
    paragraph = doc.add_paragraph()
    set_spacing(paragraph, before, after, line)
    set_font(paragraph.add_run(text), size=size, bold=bold)
    return paragraph


def add_control_marker(doc, text, centered=False, before=0, after=5):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(paragraph, before, after, 1.2)
    set_font(paragraph.add_run(text), size=11, bold=True, color=CONTROL)
    return paragraph


def add_control_field(doc, label, value, after=4):
    paragraph = doc.add_paragraph()
    set_spacing(paragraph, after=after, line=1.3)
    set_font(paragraph.add_run(f"{label}："), size=11, bold=True, color=CONTROL)
    set_font(paragraph.add_run(value), size=11)
    return paragraph


def add_image(doc, path, caption_label, caption, width_cm=10.8):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(paragraph, before=5, after=3, line=1)
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(caption_paragraph, after=8, line=1.2)
    set_font(caption_paragraph.add_run(f"{caption_label}："), size=9, bold=True, color=CONTROL)
    set_font(caption_paragraph.add_run(caption), size=9, color=MUTED)


def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    set_spacing(paragraph, before=12, after=5, line=1.2)
    set_font(paragraph.add_run(text), size=11, bold=True, color=INK)
    return paragraph


def build():
    if not HERO_IMAGE.exists() or not BODY_IMAGE.exists():
        raise FileNotFoundError("示範圖片不存在")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.4

    add_control_field(doc, "文章分類", "觀行錄")
    add_control_field(doc, "文章標題", "在一呼一吸之間：重新遇見自己")
    add_control_field(doc, "作者", "示範作者／編輯部", after=7)

    add_image(
        doc,
        HERO_IMAGE,
        "開頭圖片來源",
        "專案既有圖片（僅供範本展示）",
        width_cm=10.8,
    )

    add_control_marker(doc, "【正文開始】", centered=True, before=7, after=10)

    add_text(
        doc,
        "清晨沿著山徑緩緩往上走，風從樹梢穿過，腳下的土仍帶著昨夜的濕氣。"
        "平常總急著抵達目的地，這一天卻忽然想慢下來，聽一聽呼吸如何在身體裡進出。",
    )
    add_text(
        doc,
        "當注意力回到呼吸，肩膀自然鬆開，腳步也不再追趕。外在景物沒有改變，"
        "但心裡多出一點空間，能夠容納眼前的聲音、光線與自己的感受。",
    )

    add_heading(doc, "先安住身體，再整理心緒")
    add_text(
        doc,
        "練習不需要追求特殊經驗。先感覺雙腳與地面的接觸，再感覺腰背如何承接上半身的重量。"
        "呼吸變深或變淺都沒有關係，只要如實知道當下的狀態。",
    )

    quote_start = add_control_marker(doc, "【引述古文開始】", before=5, after=3)
    quote_start.paragraph_format.keep_with_next = True
    quote_first = add_text(
        doc,
        "《道德經》：「致虛極，守靜篤。萬物並作，吾以觀復。」",
        size=12,
        line=1.5,
        after=3,
    )
    quote_first.paragraph_format.keep_with_next = True
    quote_second = add_text(
        doc,
        "天地所以能長且久者，以其不自生，故能長生。",
        size=12,
        line=1.5,
        after=3,
    )
    quote_second.paragraph_format.keep_with_next = True
    add_control_marker(doc, "【引述古文結束】", after=7)

    add_text(
        doc,
        "所謂「觀復」，不是離開日常生活，而是在紛亂之中看見自己如何起心動念。"
        "一旦看見，便多了一個不急著反應的選擇。",
    )

    add_heading(doc, "把練習帶回人群之中")
    add_text(
        doc,
        "真正的練習往往發生在人與人相處時。有人靠近，有人離開；有人理解，也有人誤會。"
        "身體一緊，就先停一下，讓呼吸替自己留出回應的時間。",
    )

    add_image(
        doc,
        BODY_IMAGE,
        "圖片標題",
        "相聚時練習傾聽，也練習自在地成為自己",
        width_cm=10.8,
    )

    add_text(
        doc,
        "回到群體，不代表失去自己。當每個人都能穩穩站在自己的位置上，關係反而有了流動的可能。"
        "我們不必急著證明什麼，只需清楚地感受、真誠地表達，也尊重彼此不同的節奏。",
    )

    add_heading(doc, "每天留一小段空白")
    add_text(
        doc,
        "每天不妨留五分鐘，不看訊息、不處理事情，只感覺呼吸與身體。"
        "這五分鐘不為了變得更有效率，而是提醒自己：在所有角色與任務之前，我仍是一個正在呼吸的人。",
    )
    add_text(
        doc,
        "當這份覺察逐漸成為習慣，山林不再只是遠方的風景，安定也不必等待特定時刻才會出現。"
        "一呼一吸之間，我們隨時可以重新出發。",
    )

    core = doc.core_properties
    core.title = "氣機導引電子報文章 Word 範本－編輯小組版"
    core.subject = "可直接撰寫與歸檔的網站文章範本"
    core.author = "氣機導引電子報編輯部"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)

    verify = Document(OUTPUT)
    paragraphs = [paragraph for paragraph in verify.paragraphs if paragraph.text.strip()]
    texts = [paragraph.text.strip() for paragraph in paragraphs]
    headings = {
        "先安住身體，再整理心緒",
        "把練習帶回人群之中",
        "每天留一小段空白",
    }
    assert "【正文開始】" in texts
    assert "【引述古文開始】" in texts
    assert "【引述古文結束】" in texts
    assert not any(text.startswith("期數：") for text in texts)
    assert not any(text.startswith("日期：") for text in texts)
    for heading in headings:
        paragraph = next(item for item in paragraphs if item.text.strip() == heading)
        assert all(run.bold for run in paragraph.runs if run.text.strip())
        assert all(run.font.color.rgb == INK for run in paragraph.runs if run.text.strip())
    assert len(verify.inline_shapes) == 2
    image_widths = [shape.width for shape in verify.inline_shapes]
    assert image_widths[0] == image_widths[1]
    assert all(
        not paragraph.text.strip()
        for item in verify.sections
        for paragraph in item.footer.paragraphs
    )
    print(f"{OUTPUT}\nparagraphs={len(paragraphs)} headings={len(headings)} images={len(verify.inline_shapes)}")


if __name__ == "__main__":
    build()
